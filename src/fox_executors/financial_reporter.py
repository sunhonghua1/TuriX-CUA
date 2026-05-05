import pandas as pd
import pathlib
import sys
import os
import json
from openai import OpenAI
from docx import Document

def _expand_path(p):
    return str(pathlib.Path(p).expanduser().resolve())

def generate_report(excel_path, output_path):
    try:
        # 1. 自动探测输入文件（处理文件名空格问题）
        if not pathlib.Path(excel_path).exists():
            desktop = pathlib.Path.home() / "Desktop"
            possible_files = list(desktop.glob("*发票汇总*.xlsx"))
            if possible_files:
                excel_path = str(possible_files[0])
                print(f"[financial_reporter] 🔍 自动找到输入文件: {excel_path}")
            else:
                raise FileNotFoundError(f"在桌面上没找到包含'发票汇总'的 Excel 文件")

        # 2. 读取数据
        df = pd.read_excel(excel_path)
        
        # 暴力清理金额列：只保留数字和小数点
        def clean_amount(val):
            if pd.isna(val): return 0.0
            import re
            s = str(val).replace(',', '') # 先去掉千分位逗号
            nums = re.findall(r"[-+]?\d*\.\d+|\d+", s)
            return float(nums[0]) if nums else 0.0

        amounts = df['价税合计'].apply(clean_amount) if '价税合计' in df.columns else pd.Series([0.0]*len(df))
        
        summary_stats = {
            "总金额": amounts.sum(),
            "发票总数": len(df),
            "最大单笔支出": amounts.max(),
            "频繁往来方": df['销售方'].value_counts().idxmax() if '销售方' in df.columns else "未知"
        }
        
        # 2. 调用 26B 模型进行深度分析
        _base_url = os.environ.get("FOX_LOCAL_LLM_BASE_URL", "http://127.0.0.1:8000/v1")
        _model = os.environ.get("FOX_LOCAL_LLM_MODEL", "gemma-4-26b")
        client = OpenAI(base_url=_base_url, api_key="not-needed")
        
        prompt = f"""
你是一名资深财务分析师。请根据以下提取的发票汇总数据，写一份专业、简练的财务分析报告摘要。
数据：
- 发票总数：{summary_stats['发票总数']} 张
- 总报销金额：{summary_stats['总金额']:.2f} 元
- 最大单笔金额：{summary_stats['最大单笔支出']:.2f} 元
- 主要供应商：{summary_stats['频繁往来方']}

要求：
1. 评价报销频次和金额的合理性。
2. 给出 1-2 条财务优化建议。
3. 语气专业，控制在 200 字以内。
"""
        response = client.chat.completions.create(
            model=_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        analysis_text = response.choices[0].message.content.strip()

        # 3. 生成 Word 文档
        doc = Document()
        
        # 设置全局中文字体（微软雅黑）
        from docx.shared import Pt
        from docx.oxml.ns import qn
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        doc.add_heading('月度财务分析报告', 0)
        doc.add_paragraph(f'报告生成日期：{pd.Timestamp.now().strftime("%Y-%m-%d")}')
        
        doc.add_heading('一、数据概览', level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'
        
        for k, v in summary_stats.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(k)
            row_cells[1].text = str(v)

        doc.add_heading('二、专家点评 (Gemma 26B Power)', level=1)
        doc.add_paragraph(analysis_text)
        
        doc.save(output_path)
        print(f"✅ 报告已成功生成并保存至: {output_path}")
        return True
    except Exception as e:
        print(f"❌ 生成报告失败: {e}", file=sys.stderr)
        return False

if __name__ == "__main__":
    excel = _expand_path(sys.argv[1]) if len(sys.argv) > 1 else _expand_path("~/Desktop/05月发票汇总.xlsx")
    output = _expand_path(sys.argv[2]) if len(sys.argv) > 2 else _expand_path("~/Desktop/财务分析报告.docx")
    generate_report(excel, output)
