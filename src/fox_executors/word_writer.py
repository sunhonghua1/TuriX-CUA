#!/usr/bin/env python3
"""
Word 文档静默生成器 — 接收 JSON 数据，后台生成排版精美的 .docx 文件。
绝对不弹窗、不抢焦点。

用法:
    python word_writer.py ~/Desktop/报告.docx '{"title":"...","sections":[...]}'
    python word_writer.py ~/Desktop/报告.docx -  < data.json
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("[word_writer] ❌ 缺少 python-docx 库，请运行: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def run(output_path: str, data_json: str) -> int:
    # ── 解析路径 ──
    try:
        target = Path(output_path).expanduser().resolve()
    except Exception as e:
        print(f"[word_writer] ❌ 路径解析失败: {e}", file=sys.stderr)
        return 1

    # ── 目录自动创建 ──
    target.parent.mkdir(parents=True, exist_ok=True)

    # ── 解析 JSON ──
    try:
        data = json.loads(data_json)
    except json.JSONDecodeError as e:
        print(f"[word_writer] ❌ JSON 解析失败: {e}", file=sys.stderr)
        return 1

    title = data.get("title", "")
    subtitle = data.get("subtitle", "")
    sections = data.get("sections", [])

    if not title and not sections:
        print("[word_writer] ❌ 数据为空（无 title 且无 sections）", file=sys.stderr)
        return 1

    # ── 生成文档 ──
    try:
        doc = Document()

        # 大标题 — 居中
        if title:
            p_title = doc.add_heading(title, level=0)
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 副标题 — 居中
        if subtitle:
            p_sub = doc.add_paragraph(subtitle)
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p_sub.runs:
                run.italic = True

        # 章节
        for sec in sections:
            heading = sec.get("heading", "")
            body = sec.get("body", "")

            if heading:
                doc.add_heading(heading, level=1)

            if body:
                doc.add_paragraph(body)

        doc.save(str(target))

    except Exception as e:
        print(f"[word_writer] ❌ 文档生成失败: {e}", file=sys.stderr)
        return 1

    print(f"[word_writer] ✅ 报告已生成 -> {target}")
    return 0


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python word_writer.py <output.docx> <data_json_or_dash>", file=sys.stderr)
        print('  data_json: \'{"title":"...","sections":[...]}\'', file=sys.stderr)
        print("  用 - 从 stdin 读取 JSON", file=sys.stderr)
        sys.exit(1)

    output_path = sys.argv[1]
    raw = sys.argv[2]

    if raw == "-":
        data_json = sys.stdin.read()
    else:
        data_json = raw

    sys.exit(run(output_path, data_json))
